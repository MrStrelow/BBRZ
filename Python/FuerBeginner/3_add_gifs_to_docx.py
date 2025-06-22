import glob
import re
from docx import Document
from docx.shared import Inches
import os
import time

def replace_img_tags_in_docx(docx_path):
    """
    Opens a .docx file, finds <img ...> tags in paragraphs,
    and replaces them with the actual images.
    """
    try:
        document = Document(docx_path)
        print(f"-> Processing '{os.path.basename(docx_path)}'")
    except Exception as e:
        print(f"  [ERROR] Could not open file {docx_path}. Skipping. Error: {e}")
        return

    img_regex = re.compile(r'<img .*?src="([^"]+)".*?>')
    
    paragraphs_to_process = []
    for p in document.paragraphs:
        if '<img' in p.text and 'src=' in p.text:
            paragraphs_to_process.append(p)

    if not paragraphs_to_process:
        print("  -> No image tags found to replace.")
        # Even if no images, we should still rename the file
    
    for p in paragraphs_to_process:
        for match in img_regex.finditer(p.text):
            placeholder_tag = match.group(0)
            image_path = match.group(1)

            docx_dir = os.path.dirname(docx_path)
            full_image_path = os.path.join(docx_dir, image_path)

            print(f"  -> Found placeholder: {placeholder_tag}")
            print(f"     Attempting to insert image from: {full_image_path}")

            for run in p.runs:
                if placeholder_tag in run.text:
                    # Clear the placeholder text from the run
                    run.text = run.text.replace(placeholder_tag, "")
                    try:
                        # Add the picture to the same run
                        run.add_picture(full_image_path, width=Inches(4.5))
                        print(f"     [SUCCESS] Image inserted.")
                    except FileNotFoundError:
                        print(f"     [ERROR] Image not found at '{full_image_path}'.")
                    except Exception as e:
                        print(f"     [ERROR] Could not add picture. {e}")

    # --- MODIFIED SAVE LOGIC ---
    # Create the final filename by removing '-from-md' from the original name
    final_path = docx_path.replace('-from-md', '')
    document.save(final_path)
    print(f"-> Saved final version as '{os.path.basename(final_path)}'")


if __name__ == '__main__':
    DELETE_INPUT = True
    if DELETE_INPUT:
        print("\033[31mWILL DELETE input docx Files.\033[0m")
        time.sleep(3)

    print("--- Starting Post-Processing of DOCX Files ---")
    
    # --- MODIFIED SEARCH PATTERN ---
    # Use glob to find only .docx files that contain '-from-md'.
    # This automatically prevents re-processing files.
    search_pattern = './**/*-from-md.docx'
    
    files_found = glob.glob(search_pattern, recursive=True)

    if not files_found:
        print("No '...-from-md.docx' files found to process.")
    else:
        for docx_file in files_found:
            replace_img_tags_in_docx(docx_file)

            if DELETE_INPUT:
                try:
                    os.remove(docx_file)
                    print(f"-> Deleted temporary file '{os.path.basename(docx_file)}'")
                except OSError as e:
                    print(f"  [ERROR] Could not delete temporary file {docx_file}. Error: {e}")


    print("\n--- Post-Processing Complete ---")